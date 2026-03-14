"""Document processing service for parsing PDF, DOCX, XLSX files
Enhanced with AI-driven table extraction using multi-modal vision model"""
from typing import List, Dict, Tuple, Optional
from pathlib import Path
import logging
import hashlib
import re
import io

# Document parsing libraries
import pypdf
from docx import Document as DocxDocument
import openpyxl
from PIL import Image

# PDF to image conversion for vision model
try:
    import pdf2image
    HAS_PDF2IMAGE = True
except ImportError:
    HAS_PDF2IMAGE = False

from backend.config import settings
from backend.services.dashscope_service import dashscope_service

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing and parsing documents with AI-enhanced table extraction"""

    def __init__(self):
        self.chunk_size = 1000  # Characters per chunk
        self.chunk_overlap = 200  # Character overlap between chunks
        self.enable_vision_table_extraction = True  # Enable vision-based table extraction
        self.check_pdf2image()

    def check_pdf2image(self):
        """Check if pdf2image is available for vision-based table extraction"""
        if not HAS_PDF2IMAGE:
            logger.warning(
                "pdf2image not installed. Vision-based table extraction will be disabled. "
                "Install with: pip install pdf2image"
            )
            self.enable_vision_table_extraction = False
            return

        # Check if poppler is available (required for pdf2image)
        try:
            # Try to get poppler path from environment
            import shutil
            poppler_path = shutil.which('pdftoppm')
            if poppler_path:
                logger.info(f"✅ Poppler found at: {poppler_path}")
                self.enable_vision_table_extraction = True
            else:
                logger.warning(
                    "⚠️  Poppler not found (required by pdf2image on Windows). "
                    "Vision-based table extraction disabled. "
                    "Install poppler from: https://github.com/oschwartz10612/poppler-windows/"
                )
                self.enable_vision_table_extraction = False
        except Exception as e:
            logger.warning(f"⚠️  Could not verify poppler: {e}")
            self.enable_vision_table_extraction = False

    def _get_file_hash(self, file_path: str) -> str:
        """Calculate MD5 hash of file"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

    def _split_text_into_chunks(self, text: str) -> List[str]:
        """
        Split text into chunks for embedding and vector storage

        Args:
            text: Input text

        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = start + self.chunk_size

            # Try to break at word boundary
            if end < text_length:
                while end > start + self.chunk_size * 0.7:
                    if text[end] in [' ', '\n', '.', ',', ';', '!']:
                        end += 1
                        break
                    end -= 1

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - self.chunk_overlap

        return chunks

    def _convert_pdf_page_to_image(self, page, page_num: int) -> Optional[bytes]:
        """
        Convert PDF page to image for vision model analysis

        Args:
            page: pypdf Page object
            page_num: Page number

        Returns:
            Image bytes or None
        """
        if not HAS_PDF2IMAGE:
            return None

        try:
            # This would require saving PDF to temp file and using pdf2image
            # For now, we'll skip this and rely on text-based detection
            logger.info(f"Vision extraction for page {page_num}: pdf2image available but requires full PDF file")
            return None
        except Exception as e:
            logger.error(f"Error converting PDF page to image: {e}")
            return None

    def _detect_tables_in_text(self, text: str) -> bool:
        """
        Detect if text contains table-like patterns

        Args:
            text: Text to analyze

        Returns:
            True if table patterns detected
        """
        # Patterns indicating tables
        table_indicators = [
            r'\|.*\|',  # Pipe-separated columns
            r'\t.*\t',  # Tab-separated
            r'\b\d+[\.\,]\d{2}\b.*\b\d+[\.\,]\d{2}\b',  # Multiple numbers with decimals
            r'\bJan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec\b',  # Month headers
            r'\bQ1|Q2|Q3|Q4\b',  # Quarter headers
            r'Row\s+\d+',  # Row indicators
            r'Column\s+\d+',  # Column indicators
        ]

        count = sum(1 for pattern in table_indicators if re.search(pattern, text, re.IGNORECASE))
        return count >= 2

    def _extract_tables_with_vision(self, file_path: str, page_num: int, page_text: str) -> List[Dict]:
        """
        Use multi-modal vision model to extract tables from PDF page

        Args:
            file_path: Path to PDF file
            page_num: Page number
            page_text: Text content of the page

        Returns:
            List of extracted table chunks with metadata
        """
        if not self.enable_vision_table_extraction or not HAS_PDF2IMAGE:
            return []

        try:
            # Convert PDF to images using pdf2image
            # Note: This requires poppler to be installed on the system
            from pdf2image import convert_from_path

            # Convert only the specific page
            images = convert_from_path(
                file_path,
                first_page=page_num,
                last_page=page_num,
                dpi=300  # Higher DPI for better table recognition
            )

            if not images:
                logger.warning(f"No image generated for page {page_num}")
                return []

            # Get the first (and only) image
            image = images[0]

            # Convert to bytes for API call
            img_byte_arr = io.BytesIO()
            image.save(img_byte_arr, format='PNG')
            img_bytes = img_byte_arr.getvalue()

            logger.info(f"Converted page {page_num} to image: {len(img_bytes)} bytes")

            # Call vision model for table extraction
            table_result = dashscope_service.extract_table_from_image(
                img_bytes,
                page_context=page_text
            )

            if table_result and table_result.get('extracted'):
                table_chunks = []
                table_data = table_result.get('data', {})

                # Format table as text chunk
                table_text_parts = []
                if table_data.get('title'):
                    table_text_parts.append(f"Title: {table_data['title']}")

                if table_data.get('headers'):
                    table_text_parts.append(f"Headers: {' | '.join(table_data['headers'])}")

                if table_data.get('rows'):
                    table_text_parts.append("Data:")
                    for idx, row in enumerate(table_data['rows'], 1):
                        table_text_parts.append(f"  Row {idx}: {' | '.join(row)}")

                if table_text_parts:
                    table_text = '\n'.join(table_text_parts)
                    table_chunks.append({
                        'page_number': page_num,
                        'chunk_index': f"vision_table_{page_num}",
                        'text': table_text,
                        'type': 'table',
                        'ai_extracted': True,
                        'method': 'vision'  # Mark as vision-extracted
                    })
                    logger.info(f"✅ Vision-extracted table from page {page_num}: {len(table_data.get('rows', []))} rows")

                return table_chunks

            return []

        except ImportError:
            logger.warning("pdf2image not available for vision table extraction")
            self.enable_vision_table_extraction = False
            return []
        except Exception as e:
            logger.error(f"Vision table extraction failed for page {page_num}: {e}")
            return []

    def _extract_tables_with_ai(self, text: str, page_num: int) -> List[Dict]:
        """
        Use AI to extract structured table data from text (fallback method)

        Args:
            text: Text potentially containing tables
            page_num: Page number for metadata

        Returns:
            List of extracted table chunks with metadata
        """
        if not self._detect_tables_in_text(text):
            return []

        try:
            # 直接使用 DashScope API 而不是通过 generate_response,以避免简洁回答的 prompt 干扰
            from dashscope import Generation

            # 使用特定的表格提取 prompt,不使用简洁回答的约束
            system_prompt = """You are a table extraction expert. Your task is to extract and structure table data from text.

STRICT INSTRUCTIONS:
1. Extract ALL tables found in the text
2. Use EXACTLY this output format for each table:
   [TABLE 1]
   Title: [table title if any, or leave blank]
   Headers: [col1, col2, col3, ...]
   Row 1: [val1, val2, val3, ...]
   Row 2: [val1, val2, val3, ...
   [END TABLE 1]

3. Preserve numerical values EXACTLY as they appear
4. Include all rows from the table
5. Output ONLY the table data, no explanations or summaries

IMPORTANT: Do not simplify or summarize. Extract the complete table structure."""

            user_prompt = f"""Extract tables from the following text:

{text}

Extract all tables using the exact format specified."""

            response = Generation.call(
                model=dashscope_service.llm_model,
                messages=[
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': user_prompt}
                ],
                result_format='message',
                max_tokens=2048,
                temperature=0.1  # Low temperature for consistent formatting
            )

            # Parse AI response for tables
            table_chunks = []
            if response.status_code == 200:
                result_text = response.output.choices[0]['message']['content']

                if "[TABLE" in result_text:
                    # Split by table markers
                    table_sections = re.split(r'\[TABLE \d+\]', result_text)

                    for idx, table_section in enumerate(table_sections[1:], 1):  # Skip first empty section
                        if table_section.strip():
                            # Create structured chunk for table
                            table_chunks.append({
                                'page_number': page_num,
                                'chunk_index': f"ai_table_{idx}",
                                'text': f"[Table {idx}] {table_section}",
                                'type': 'table',
                                'ai_extracted': True,
                                'method': 'text'  # Mark as text-based AI extraction
                            })

                            logger.info(f"AI-extracted table {idx} from page {page_num} (text-based)")
                else:
                    logger.warning(f"AI response did not contain expected [TABLE] format for page {page_num}")
                    # Fallback: use the raw response as a table chunk
                    if result_text.strip():
                        table_chunks.append({
                            'page_number': page_num,
                            'chunk_index': f"ai_table_raw",
                            'text': f"[AI Extracted Table Data]\n{result_text}",
                            'type': 'table',
                            'ai_extracted': True,
                            'method': 'text'
                        })
                        logger.info(f"AI-extracted raw table data from page {page_num}")

            return table_chunks

        except Exception as e:
            logger.error(f"AI table extraction failed: {e}")
            return []

    def process_pdf(self, file_path: str) -> Tuple[str, List[Dict]]:
        """
        Process PDF file and extract text with AI-enhanced table extraction

        Args:
            file_path: Path to PDF file

        Returns:
            Tuple of (full_text, list of page chunks with metadata)
        """
        try:
            text = ""
            pages = []

            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    page_text = page.extract_text()
                    text += page_text + "\n\n"

                    # Regular text chunks
                    chunks = self._split_text_into_chunks(page_text)
                    for chunk_idx, chunk in enumerate(chunks):
                        pages.append({
                            'page_number': page_num,
                            'chunk_index': chunk_idx,
                            'text': chunk,
                            'type': 'text'
                        })

                    # Try vision-based table extraction first
                    table_chunks = []
                    if self.enable_vision_table_extraction and self._detect_tables_in_text(page_text):
                        logger.info(f"🔍 Attempting vision-based table extraction for page {page_num}")
                        table_chunks = self._extract_tables_with_vision(file_path, page_num, page_text)

                    # Fallback to text-based AI extraction if vision fails
                    if not table_chunks and self._detect_tables_in_text(page_text):
                        logger.info(f"📝 Fallback to text-based AI table extraction for page {page_num}")
                        table_chunks = self._extract_tables_with_ai(page_text, page_num)

                    pages.extend(table_chunks)

            vision_tables = len([p for p in pages if p.get('method') == 'vision'])
            ai_tables = len([p for p in pages if p.get('method') == 'text'])
            logger.info(
                f"✅ Successfully processed PDF: {len(pages)} chunks total "
                f"({vision_tables} vision-tables, {ai_tables} AI-tables)"
            )
            return text, pages

        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise

    def process_docx(self, file_path: str) -> Tuple[str, List[Dict]]:
        """
        Process DOCX file and extract text

        Args:
            file_path: Path to DOCX file

        Returns:
            Tuple of (full_text, list of chunks with metadata)
        """
        try:
            doc = DocxDocument(file_path)
            text = ""
            chunks = []

            # Extract paragraphs
            paragraph_texts = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n\n".join(paragraph_texts)

            # Split into chunks
            chunk_list = self._split_text_into_chunks(text)
            for idx, chunk in enumerate(chunk_list):
                chunks.append({
                    'page_number': 1,  # DOCX doesn't have pages
                    'chunk_index': idx,
                    'text': chunk
                })

            # Extract tables using AI if present
            if doc.tables:
                for table_idx, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)

                    # AI-enhanced table extraction
                    if table_data:
                        table_text = "\n".join([" | ".join(row) for row in table_data])
                        table_chunks = self._split_text_into_chunks(table_text)
                        for chunk_idx, chunk in enumerate(table_chunks):
                            chunks.append({
                                'page_number': 1,
                                'chunk_index': len(chunks) + chunk_idx,
                                'text': f"[Table {table_idx + 1}]\n{chunk}",
                                'type': 'table',
                                'method': 'text'
                            })

            logger.info(f"Successfully processed DOCX: {len(chunks)} chunks extracted")
            return text, chunks

        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise

    def process_xlsx(self, file_path: str) -> Tuple[str, List[Dict]]:
        """
        Process Excel file and extract data

        Args:
            file_path: Path to XLSX file

        Returns:
            Tuple of (full_text, list of chunks with metadata)
        """
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text = ""
            chunks = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]

                # Extract table data
                sheet_data = []
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_data):  # Skip empty rows
                        sheet_data.append(row_data)

                # Convert to text
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += "\n".join([" | ".join(row) for row in sheet_data])
                text += sheet_text + "\n\n"

                # AI-enhanced table processing for structured data
                if sheet_data:
                    # Create chunk with table representation
                    table_chunks = self._split_text_into_chunks(sheet_text)
                    for chunk_idx, chunk in enumerate(table_chunks):
                        chunks.append({
                            'page_number': 1,
                            'chunk_index': len(chunks) + chunk_idx,
                            'text': chunk,
                            'type': 'table',
                            'sheet_name': sheet_name,
                            'method': 'direct'
                        })

            logger.info(f"Successfully processed XLSX: {len(chunks)} chunks extracted")
            return text, chunks

        except Exception as e:
            logger.error(f"Error processing XLSX {file_path}: {e}")
            raise

    def process_document(self, file_path: str, file_format: str) -> Tuple[str, List[Dict]]:
        """
        Process document based on format

        Args:
            file_path: Path to document
            file_format: File format (pdf, docx, xlsx)

        Returns:
            Tuple of (full_text, list of chunks with metadata)
        """
        file_format = file_format.lower()

        if file_format == 'pdf':
            return self.process_pdf(file_path)
        elif file_format == 'docx':
            return self.process_docx(file_path)
        elif file_format == 'xlsx':
            return self.process_xlsx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_format}")

    def get_file_info(self, file_path: str) -> Dict:
        """
        Get file metadata

        Args:
            file_path: Path to file

        Returns:
            Dictionary with file information
        """
        path = Path(file_path)
        return {
            'name': path.stem,
            'filename': path.name,
            'size': path.stat().st_size,
            'hash': self._get_file_hash(file_path)
        }


# Global instance
document_processor = DocumentProcessor()
